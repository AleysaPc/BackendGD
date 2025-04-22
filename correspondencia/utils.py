from io import BytesIO
from django.http import HttpResponse
from docx import Document
from docx.shared import Pt
from django.utils.timezone import now
import os
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

def generar_documento_word(docSaliente):
    """Genera un documento Word a partir de un objeto CorrespondenciaSaliente."""

    doc = Document()

    # Fecha
    doc.add_paragraph(f"La Paz, {docSaliente.fecha_envio.strftime('%d-%m-%Y')}")

    # CITE en negrita
    parrafo_cite = doc.add_paragraph()
    run_cite = parrafo_cite.add_run(f"CITE: FSTL-FTA/DPTO/LP/ N° {docSaliente.cite}")
    run_cite.bold = True

    # Señor:
    doc.add_paragraph("Señor:")

    if docSaliente.correspondencia.contacto.titulo_profesional == "Ingeniero" :
        titulo = "Ing."
    elif docSaliente.correspondencia.contacto.titulo_profesional == "Licenciado":
        titulo = "Lic."
    elif docSaliente.correspondencia.contacto.titulo_profesional == "Doctor":
        titulo = "Dr."
    elif docSaliente.correspondencia.contacto.titulo_profesional == "Abogado":
        titulo = "Abog."
    elif docSaliente.correspondencia.contacto.titulo_profesional == "Profesor":
        titulo = "Prof."
    elif docSaliente.correspondencia.contacto.titulo_profesional == "Magister":
        titulo = "Mgs."
    else:
        titulo = ""
    # Nombre del contacto
    doc.add_paragraph(f"{titulo} {docSaliente.correspondencia.contacto.nombre_contacto}")

    # Apellidos del contacto
    doc.add_paragraph(f"{docSaliente.correspondencia.contacto.apellido_pat_contacto} {docSaliente.correspondencia.contacto.apellido_mat_contacto}")

    # Cargo
    doc.add_paragraph(docSaliente.correspondencia.contacto.cargo.upper())

    # Institución
    doc.add_paragraph(str(docSaliente.correspondencia.contacto.institucion).upper())

    # Presente
    doc.add_paragraph("Presente.-")

    # Referencia alineada a la derecha, subrayada y en negrita
    parrafo_ref = doc.add_paragraph()
    parrafo_ref.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
    run_ref = parrafo_ref.add_run(f"Ref.: {docSaliente.correspondencia.referencia}")
    run_ref.bold = True
    run_ref.underline = True

    # Texto de cortesía
    doc.add_paragraph("De nuestra mayor consideración:")

    # Descripción
    doc.add_paragraph(docSaliente.correspondencia.descripcion)

    # Guardar el archivo en un buffer
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)

    # Devolver el archivo como respuesta
    response = HttpResponse(
        buffer.getvalue(),
        content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
    )
    response['Content-Disposition'] = f'attachment; filename=correspondencia_{docSaliente.cite}.docx'
    
    return response
